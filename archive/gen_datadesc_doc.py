"""
Generate CDIFDataDescriptionClasses.docx from CDIFDiscoveryClasses.docx.

Strategy: Copy the source document (preserving all formatting), then make
targeted modifications using python-docx paragraph/run manipulation.
"""
import copy
import shutil
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = 'CDIFDiscoveryClasses.docx'
DST = 'CDIFDataDescriptionClasses.docx'

# Step 1: Copy the file
shutil.copy2(SRC, DST)

# Step 2: Open and modify
doc = Document(DST)
paras = doc.paragraphs


def clear_and_set_text(para, text):
    """Clear all runs in a paragraph and set new text (single run).
    Also removes hyperlinks and other inline elements that contain runs."""
    # Remove all child elements except pPr (paragraph properties)
    children_to_remove = []
    for child in para._element:
        tag = child.tag.split('}')[-1]
        if tag != 'pPr':
            children_to_remove.append(child)
    for child in children_to_remove:
        para._element.remove(child)
    # Add a single new run with the text
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    para._element.append(r)


def add_paragraph_after(para, text, style_name):
    """Insert a new paragraph immediately after the given paragraph element."""
    new_p = OxmlElement('w:p')
    # Set the style via pPr
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), style_name)
    pPr.append(pStyle)
    new_p.append(pPr)
    # Add a run with the text
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    new_p.append(r)
    # Insert after
    para._element.addnext(new_p)
    return new_p


def insert_multiple_after(anchor_element, items):
    """Insert multiple (text, style_name) pairs after anchor, in order.
    Returns the last inserted element."""
    current = anchor_element
    for text, style_name in items:
        new_p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        pStyle = OxmlElement('w:pStyle')
        # Map style names to their Word style IDs
        style_map = {
            'Heading 2': 'Heading2',
            'Heading 3': 'Heading3',
            'Heading 4': 'Heading4',
            'Heading 5': 'Heading5',
            'Normal': 'Normal',
            'property': 'property',
        }
        pStyle.set(qn('w:val'), style_map.get(style_name, style_name))
        pPr.append(pStyle)
        new_p.append(pPr)
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        new_p.append(r)
        current.addnext(new_p)
        current = new_p
    return current


# ---- Modification 1: Title (paragraph 0) ----
clear_and_set_text(paras[0], 'CDIF Data Description Classes')

# ---- Modification 2: Intro paragraph (paragraph 3) ----
new_intro = (
    'This document includes all of the classes and properties for the CDIF Data Description profile '
    'implementation using the schema.org vocabulary. The Data Description profile extends the Discovery '
    'profile with detailed variable documentation and data structure description. It composes: '
    '1) the cdifCore base properties (mandatory and optional); '
    '2) discovery-oriented properties (measurement technique, spatial/temporal coverage, quality); and '
    '3) data description extensions that add requirements for variable identifiers, physical data types, '
    'and distribution-level file characterization. '
    'Conformance to this profile entails populating all mandatory content from cdifCore, using recommended '
    'discovery properties, and providing the additional data description constraints. '
    'The implementation target is an rdf serialization, which is an open world logical model; users are '
    'thus free to add additional properties that they find useful for dataset documentation in their '
    'community, but these can be ignored by other users without penalty.'
)
clear_and_set_text(paras[3], new_intro)

# ---- Modification 3: Model intro paragraph (paragraph 6) ----
new_model_intro = (
    'The information model for the data description profile is defined in the cdifBook. '
    'This section outlines a logical model focused on a JSON-LD implementation of the content items '
    'defined in the information model. All classes and properties are implemented with schema.org types '
    'and attributes unless there is a prefix indicating use of elements from other vocabularies. '
    'See the context section for prefixes used and their mapping to URIs.'
)
clear_and_set_text(paras[6], new_model_intro)

# ---- Modification 4: distribution property (paragraph 30) ----
existing_dist = paras[30].text
new_dist = (
    existing_dist + ' At the Data Description level, distribution items gain additional properties: '
    'cdi:characterSet, cdi:fileSize, and cdi:fileSizeUofM for file characterization.'
)
clear_and_set_text(paras[30], new_dist)

# ---- Modification 5: variableMeasured description (paragraph 70) ----
new_vm = (
    '[Optional, Repeatable] (PropertyValue-(variableMeasured) with cdifVariableMeasured extensions). '
    'At the Data Description level, each variableMeasured item MUST have an @id property and is extended '
    'with DDI-CDI InstanceVariable properties for physical data types, variable roles, and cross-variable '
    'references. See PropertyValue-(variableMeasured) and CdifVariableMeasured classes below.'
)
clear_and_set_text(paras[70], new_vm)

# ---- Modification 6: conformsTo (paragraph 195) ----
new_conformsto = (
    '[Required, Repeatable] (object reference). Identifiers for conformance classes/profiles that the '
    'metadata record follows. For CDIF data description must include '
    '"https://w3id.org/cdif/discovery/1.1", "https://w3id.org/cdif/core/1.1", and '
    '"https://w3id.org/cdif/data_description/1.1" because conforms to all three profiles.'
)
clear_and_set_text(paras[195], new_conformsto)

# ---- Modification 7: Insert "Properties added in Data Description Profile" ----
# Insert after paragraph 76 (dqv:hasQualityMeasurement property text), before paragraph 77 ("Other Classes")
# We need to find the XML element for paragraph 76
anchor_76 = paras[76]._element

dd_properties = [
    ('Properties added in Data Description Profile', 'Heading 4'),
    ('cdi:characterSet', 'Heading 5'),
    ('[Optional] (string) The character set used in the distribution (e.g., UTF-8, ASCII). This property is used on distribution items.', 'property'),
    ('cdi:fileSize', 'Heading 5'),
    ('[Optional] (number) The size of the distribution file. This property is used on distribution items.', 'property'),
    ('cdi:fileSizeUofM', 'Heading 5'),
    ('[Optional] (string) Unit of measure for the file size (e.g., bytes, KB, MB, GB). This property is used on distribution items.', 'property'),
]

insert_multiple_after(anchor_76, dd_properties)

# ---- Modification 8: Insert "Classes added by CDIF Data Description profile" ----
# This should go after the Discovery classes section.
# Paragraph 216 = "Heading 2: Classes added by CDIF Discovery profile"
# The Discovery classes end at paragraph 269 (last property of DefinedTerm: inDefinedTermSet)
# Paragraph 270 = "Heading 2: Data types used for CDIF Core"
# Insert before paragraph 270

anchor_269 = paras[269]._element  # inDefinedTermSet property text

dd_classes = [
    ('Classes added by CDIF Data Description profile', 'Heading 2'),
    ('CdifVariableMeasured', 'Heading 3'),
    (
        'This class describes a CdifVariableMeasured node, extending schema.org PropertyValue '
        'with DDI-CDI InstanceVariable properties. It composes the base variableMeasured building block via allOf. '
        'At the Data Description level, a variableMeasured item uses this extended class to provide '
        'physical data type information, variable roles, and cross-variable references.',
        'Normal'
    ),
    ('type', 'Heading 5'),
    ('[Required, Repeatable] (string.uri) Must include cdi:InstanceVariable. Additional types may be included.', 'property'),
    ('@id', 'Heading 5'),
    (
        '[Required] (string) URI identifier for this variable. Required at Data Description level so '
        'physical mappings can reference this variable via cdi:formats_InstanceVariable.',
        'property'
    ),
    ('cdi:identifier', 'Heading 5'),
    ('[Optional] (string) DDI-CDI identifier for this variable.', 'property'),
    ('cdi:physicalDataType', 'Heading 5'),
    (
        '[Optional, Repeatable] (string, object reference, or DefinedTerm) Identifier or name for the '
        'physical data type concept. Links to controlled vocabularies for data types.',
        'property'
    ),
    ('cdi:intendedDataType', 'Heading 5'),
    (
        '[Optional] (string) The intended data type for values of this variable, from DDI-CDI '
        'RepresentedVariable.hasIntendedDataType. Recommended values are XML Schema datatypes '
        '(e.g. https://www.w3.org/TR/xmlschema-2/#decimal).',
        'property'
    ),
    ('cdi:role', 'Heading 5'),
    (
        '[Optional] (string, enum) Specifies the relation of the variable to the data structure, '
        'corresponding to DDI-CDI DataStructureComponent subtypes. Values: MeasureComponent, '
        'AttributeComponent, DimensionComponent, DescriptorComponent, ReferenceValueComponent.',
        'property'
    ),
    ('cdi:describedUnitOfMeasure', 'Heading 5'),
    (
        '[Optional] (DefinedTerm) A structured unit of measure from a controlled vocabulary, from '
        'DDI-CDI RepresentedVariable.describedUnitOfMeasure.',
        'property'
    ),
    ('cdi:simpleUnitOfMeasure', 'Heading 5'),
    ('[Optional] (string, object reference, or DefinedTerm) Simple unit reference.', 'property'),
    ('cdi:uses', 'Heading 5'),
    (
        '[Optional, Repeatable] (string, object reference, or DefinedTerm) Essentially the same as '
        'schema:propertyID. References to concepts that this variable measures or represents.',
        'property'
    ),
    ('cdi:name', 'Heading 5'),
    ('[Optional] (string) DDI-CDI Concept.name. The name of this variable in the DDI-CDI model.', 'property'),
    ('cdi:displayLabel', 'Heading 5'),
    ('[Optional] (string) DDI-CDI Concept.displayLabel. A human-readable label for display purposes.', 'property'),
    ('cdi:qualifies', 'Heading 5'),
    ('[Optional] (object with @id) Reference to another instance variable defined for this dataset.', 'property'),
]

insert_multiple_after(anchor_269, dd_classes)

# ---- Modification 9: Update conformsTo in example code (paragraph 418) ----
# Paragraph 418 has: "dcterms:conformsTo":{"@id":"https://w3id.org/cdif/discovery/1.1"}
# Update to show data_description conformance
if 'conformsTo' in paras[418].text:
    clear_and_set_text(paras[418],
        '"dcterms:conformsTo":[{"@id":"https://w3id.org/cdif/core/1.1"},{"@id":"https://w3id.org/cdif/discovery/1.1"},{"@id":"https://w3id.org/cdif/data_description/1.1"}]'
    )

# Save
doc.save(DST)
print(f'Generated {DST} successfully.')
